from django.utils.html import strip_tags

from docx.shared import Inches
from docx.enum.section import WD_ORIENT

from utils.helper import DOCXReport


class SpanFactory(object):

    def __init__(self):
        self.obj = []

    def append(self, text, bold=False, italic=False, new_line=False):
        if new_line:
            text += u"\n"

        self.obj.append( {
            "text": text,
            "bold": bold,
            "italic": italic
        })

    def get_spans(self):
        return self.obj


def getDoses(ag):
    # dictionary-mapping where keys are dose-units and doses are list of floats
    doses = {}
    for d in ag['dosing_regime']['doses']:
        units = d['dose_units']['units']
        if units not in doses:
            doses[units] = []
        doses[units].append(d['dose'])
    return doses

def getDoseText(doses):
    dose_txt = []
    for unit in doses.keys():
        dose_txt.append(u"{0} {1}".format(u", ".join([unicode(dose) for dose in doses[unit]]), unit))
    return u"\n".join(dose_txt)

def getPurityText(exp):
    if exp["purity_available"]:
        return ">{0}%".format(exp["purity"])
    else:
        return "NR"

def getCOItext(study):
    txt = study["coi_reported"]
    if study["coi_details"] != "":
        txt = u"{} ({})".format(txt, study["coi_details"])
    return txt

def getAdditionalEndpoints(eps):
    return u", ".join(ep["name"] for ep in eps if ep["data_extracted"]==False)

def getEndpointsText(eps):
    return u", ".join(ep["name"] for ep in eps if ep["data_extracted"]==True)

def getNText(eps):
    ns = []
    for ep in eps:
        for eg in ep["endpoint_group"]:
            ns.append(eg["n"])
    if len(ns)==0:
        return "Not available"
    nmin = min(ns)
    nmax = max(ns)
    if nmin==nmax:
        return unicode(nmin)
    else:
        return u"{}-{}".format(nmin, nmax)

def getStatisticalAnalysis(eps):
    return u"; ".join(set([ep["statistical_test"] for ep in eps ]))

def getStatisticalPowers(eps):
    return u"; ".join(set([ep["power_notes"] for ep in eps ]))

class EndpointDOCXReport(DOCXReport):

    def printField(self, spans, heading, text):
        spans.append(heading, bold=True)
        spans.append(text, new_line=True)

    def getStudyDetails(self, ag, doses):
        spans = SpanFactory()

        txt = "({})\n".format(ag['experiment']['study']['short_citation'])
        spans.append(txt, bold=True)

        self.printField(spans, "Species: ", ag["species"])
        self.printField(spans, "Strain (source): ", ag["strain"] + u" (data not in HAWC)")
        self.printField(spans, "Sex: ", ag["sex"])
        self.printField(spans, "Doses: ", getDoseText(doses))
        self.printField(spans, "Purity (source): ", getPurityText(ag["experiment"]) + u" (data not in HAWC)")
        self.printField(spans, "Dosing period: ", strip_tags(ag["dosing_regime"]["description"]))
        self.printField(spans, "Route: ", ag["dosing_regime"]["route_of_exposure"])
        self.printField(spans, "Diet: ", "(data not in HAWC)")
        self.printField(spans, "Negative controls: ", ag["dosing_regime"]["negative_control"])
        self.printField(spans, "Funding source: ", ag["experiment"]["study"]["funding_source"])
        self.printField(spans, "Author conflict of interest: ", getCOItext(ag["experiment"]["study"]))
        self.printField(spans, "Comments: ", strip_tags(ag["experiment"]["description"]))

        return spans.get_spans()

    def getHealthOutcomes(self, ag):
        spans = SpanFactory()
        self.printField(spans, "Endpoints: ", getEndpointsText(ag["eps"]))
        self.printField(spans, "Age at exposure: ", ag["lifestage_exposed"])
        self.printField(spans, "Age at assessment: ", ag["lifestage_assessed"])
        self.printField(spans, "N: ", getNText(ag["eps"]))
        self.printField(spans, "Statistical analysis: ", getStatisticalAnalysis(ag["eps"]))
        self.printField(spans, "Control for litter effects: ", " (data not in HAWC)")
        self.printField(spans, "Statistical power: ", getStatisticalPowers(ag["eps"]))
        self.printField(spans, "Additional endpoints not extracted: ", getAdditionalEndpoints(ag["eps"]))
        return spans.get_spans()

    def build_summary_table(self, ag):

        doses = getDoses(ag)
        firstDose = doses.keys()[0]

        # build header rows
        cells = [
            {"row": 0, "col": 0, "text": r'Template Option 1: Animal Study', "colspan": 6},
            {"row": 1, "col": 0, "text": r'Reference, Animal Model, and Dosing'},
            {"row": 1, "col": 1, "text": r'Health Outcome'},
            {"row": 1, "col": 2, "text": r'Results', "colspan": 4},
            {"row": 2, "col": 2, "text": u'Dose ({0})'.format(firstDose)},
            {"row": 2, "col": 3, "text": r'N'},
            {"row": 2, "col": 4, "text": u'Mean \u00B1 SD'},
            {"row": 2, "col": 5, "text": r'% control (95% CI)'},
        ]

        # build summary rows
        ref = {"row":2, "col":0, "runs": self.getStudyDetails(ag, doses)}
        ho = {"row":2, "col":1, "runs": self.getHealthOutcomes(ag)}

        # build endpoint rows
        rows = 3
        for ep in ag["eps"]:

            if ep["data_extracted"] == False:
                continue

            cells.append({"row": rows, "col": 2, "text": ep['name'], "colspan": 4})
            data_type = ep["data_type"]
            rows +=1
            for i, eg in enumerate(ep["endpoint_group"]):
                cells.append({"row": rows, "col": 2, "text": unicode(doses[firstDose][i])})
                cells.append({"row": rows, "col": 3, "text": unicode(eg["n"])})

                col1 = ""
                col2 = ""
                if data_type == "C":
                    col1 = u'{0} \u00B1 {1:.2f}'.format(eg["response"], eg["stdev"])
                    col2 = u'{0:.1f} ({1:.1f}, {2:.1f})'.format(eg["percentControlMean"], eg["percentControlLow"], eg["percentControlHigh"])
                elif data_type in ["D", "DC"]:
                    col1 = u'{0}/{1}'.format(eg["incidence"], eg["n"])
                    col2 = u'NR'
                else:  # ["P"]
                    col1 = u'NR'
                    col2 = u'{0:.1f}'.format(eg["percentControlMean"])
                    if eg["percentControlLow"]:
                        col2 += u"({0:.1f}, {1:.1f})".format(eg["percentControlLow"], eg["percentControlHigh"])

                cells.append({"row": rows, "col": 4, "text": col1})
                cells.append({"row": rows, "col": 5, "text": col2})
                rows +=1

        # adjust rowspans for summary rows
        ref["rowspan"] = rows-2
        ho["rowspan"] = rows-2
        cells.extend([ref, ho])

        # send to the printers...
        self.build_table(rows, 6, cells)
        self.doc.add_paragraph("\n")

    def apply_context(self):
        doc = self.doc
        d = self.context

        # make landscape
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width  = Inches(11)
        section.page_height  = Inches(8.5)

        # title
        txt = "Bioassay report: ".format(d['assessment']['name'])
        doc.add_heading(txt, 0)

        # loop through each study
        for study in d['studies']:
            doc.add_heading(study['short_citation'], 1)
            for exp in study['exps']:
                for ag in exp['ags']:
                    self.build_summary_table(ag)

